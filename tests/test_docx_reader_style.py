from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from core.parser import DOCXReader


def test_docx_reader_captures_style_bold_alignment_font_and_uppercase_ratio(tmp_path):
    path = tmp_path / "style-aware.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Heading 1"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("NAS PEMBIMBING")
    run.bold = True
    run.font.size = Pt(18)
    document.save(path)

    block = DOCXReader().read(str(path))[0]

    assert block.style_name == "Heading 1"
    assert block.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert block.has_bold is True
    assert block.max_font_size == 18
    assert block.uppercase_ratio == 1.0
