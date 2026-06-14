import io
import logging
from pathlib import Path
from typing import Protocol

from core.raw_block import RawBlock, uppercase_ratio


logger = logging.getLogger(__name__)


class PDFTextExtractionError(RuntimeError):
    pass


class DocumentReader(Protocol):
    def read(self, file_path: str) -> list[RawBlock]:
        ...


class DOCXReader:
    source_type = "docx"

    def read(self, file_path: str) -> list[RawBlock]:
        try:
            import docx
        except ImportError:
            return [RawBlock("Error: python-docx not installed.", "Error", 0)]

        document = docx.Document(file_path)
        blocks: list[RawBlock] = []
        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.rstrip()
            if text.strip():
                font_sizes = [
                    float(run.font.size.pt)
                    for run in paragraph.runs
                    if run.font.size is not None
                ]
                blocks.append(
                    RawBlock(
                        text=text,
                        style_name=paragraph.style.name,
                        index=index,
                        alignment=paragraph.alignment,
                        has_bold=any(run.bold for run in paragraph.runs),
                        max_font_size=max(font_sizes) if font_sizes else None,
                        uppercase_ratio=uppercase_ratio(text),
                        source_type=self.source_type,
                        paragraph_index=index,
                        style=paragraph.style.name,
                    )
                )
        logger.info("DOCXReader read %s block(s) from %s", len(blocks), file_path)
        return blocks


class OCRReader:
    source_type = "pdf"

    def read(self, file_path: str) -> list[RawBlock]:
        try:
            import fitz
            import pytesseract
            from PIL import Image
        except ImportError as exc:
            raise RuntimeError(
                "OCR membutuhkan PyMuPDF, Pillow, dan pytesseract. "
                "Pastikan dependency dan aplikasi Tesseract OCR tersedia."
            ) from exc

        blocks: list[RawBlock] = []
        logger.info("OCRReader started for %s", file_path)
        document = fitz.open(file_path)
        for page_index, page in enumerate(document, start=1):
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            image = Image.open(io.BytesIO(pixmap.tobytes("png")))
            text = pytesseract.image_to_string(image).strip()
            confidence = None
            try:
                data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                values = [
                    float(value)
                    for value in data.get("conf", [])
                    if str(value).strip() not in {"", "-1"}
                ]
                if values:
                    confidence = sum(values) / len(values)
            except Exception:
                confidence = None
            if text:
                blocks.append(
                    RawBlock(
                        text=text,
                        source_type=self.source_type,
                        page_number=page_index,
                        paragraph_index=0,
                        metadata={"ocr_used": True, "confidence": confidence},
                    )
                )
        logger.info("OCRReader produced %s block(s) from %s", len(blocks), file_path)
        return blocks


class PDFReader:
    source_type = "pdf"

    def __init__(self, use_ocr: bool = False, ocr_reader: OCRReader | None = None) -> None:
        self.use_ocr = use_ocr
        self.ocr_reader = ocr_reader or OCRReader()
        self.warnings: list[str] = []

    def read(self, file_path: str) -> list[RawBlock]:
        try:
            import fitz
        except ImportError as exc:
            raise RuntimeError("Import PDF membutuhkan PyMuPDF. Install dependency `PyMuPDF`.") from exc

        logger.info("PDFReader started for %s (use_ocr=%s)", file_path, self.use_ocr)
        document = fitz.open(file_path)
        blocks: list[RawBlock] = []
        for page_index, page in enumerate(document, start=1):
            text = (page.get_text("text") or "").strip()
            if text:
                blocks.append(
                    RawBlock(
                        text=text,
                        source_type=self.source_type,
                        page_number=page_index,
                        paragraph_index=0,
                        metadata={"ocr_used": False},
                    )
                )

        if blocks:
            logger.info("PDFReader extracted %s text block(s) from %s", len(blocks), file_path)
            return blocks

        message = "PDF ini kemungkinan hasil scan/gambar. Gunakan OCR atau upload DOCX."
        self.warnings.append(message)
        logger.warning("PDFReader found no text in %s", file_path)
        if self.use_ocr:
            ocr_blocks = self.ocr_reader.read(file_path)
            if ocr_blocks:
                self.warnings.append("Hasil OCR mungkin tidak 100% akurat. Mohon periksa kembali teks sebelum export.")
                return ocr_blocks
        raise PDFTextExtractionError(message)


def read_document_blocks(file_path: str, use_ocr: bool = False) -> list[RawBlock]:
    suffix = Path(file_path).suffix.lower()
    if suffix == ".docx":
        return DOCXReader().read(file_path)
    if suffix == ".pdf":
        return PDFReader(use_ocr=use_ocr).read(file_path)
    raise ValueError("Format file belum didukung. Gunakan .docx atau .pdf.")

def extract_cover_title_from_docx(doc_path: str) -> str:
    from docx import Document
    doc = Document(doc_path)
    
    best_text = ""
    max_weight = 0
    
    # Evaluasi maksimal 5 paragraf pertama di halaman awal
    for p in doc.paragraphs[:5]:
        text = p.text.strip()
        if not text:
            continue
            
        weight = 0
        # Heuristik 1: Ukuran font paragraf atau run teks
        if p.style.font.size:
            weight += p.style.font.size.pt
        
        # Heuristik 2: Atribut cetak tebal (Bold)
        if p.style.font.bold:
            weight += 15
            
        # Heuristik 3: Seluruh teks menggunakan HURUF KAPITAL
        if text.isupper():
            weight += 10
            
        if weight > max_weight:
            max_weight = weight
            best_text = text
            
    return best_text if best_text else "TATA IBADAH REKAYASA"

def extract_cover_title_from_pdf(pdf_path: str) -> str:
    try:
        import fitz
    except ImportError:
        return "TATA IBADAH"
        
    doc = fitz.open(pdf_path)
    if len(doc) == 0:
        return ""
        
    page = doc[0] # Ambil halaman pertama saja
    blocks = page.get_text("dict").get("blocks", [])
    
    best_text = ""
    max_font_size = 0
    
    for b in blocks:
        if "lines" in b:
            for l in b["lines"]:
                for s in l["spans"]:
                    text = s.get("text", "").strip()
                    size = s.get("size", 0)
                    
                    # Heuristik: Cari ukuran font murni terbesar di page 1
                    if size > max_font_size and len(text) > 3:
                        max_font_size = size
                        best_text = text
                        
    return best_text if best_text else "TATA IBADAH"

